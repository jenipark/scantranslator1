# app.py
import streamlit as st
import google.genai as genai
from google.genai.errors import APIError
from google.genai.types import Part
from PIL import Image
import io, os, json, re, hashlib
from dotenv import load_dotenv

# PDF & export helpers
import fitz  # PyMuPDF
import pandas as pd
from docx import Document
from docx.shared import Pt
from datetime import datetime

# ─────────────────────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────────────────────
st.set_page_config(page_title="ScanTranslate", layout="wide")

# ─────────────────────────────────────────────────────────────
# ENV / CLIENT
# ─────────────────────────────────────────────────────────────
load_dotenv()
API_KEY = os.getenv("GEMINI_API_KEY")

client = None
if API_KEY:
    try:
        client = genai.Client(api_key=API_KEY)
    except Exception as e:
        st.error(f"⚠️ Gemini 클라이언트 초기화 오류: {e}")
        client = None
else:
    st.warning(
        "⚠️ **GEMINI_API_KEY** 환경 변수가 설정되지 않았거나 `.env` 파일 로드에 실패했습니다. "
        "OCR/번역 기능이 작동하지 않습니다."
    )

# ─────────────────────────────────────────────────────────────
# I18N
# ─────────────────────────────────────────────────────────────
TEXTS = {
    "ko": {
        "title_main": "ScanTranslate: 한국어 → 필리핀어 OCR 도구",
        "title_sub": "이미지·PDF 텍스트를 추출/번역하고 편집, 복사, 내보내기까지 가능합니다.",
        "app_language": "앱 표시 언어",
        "extract_header": "텍스트 추출 및 번역",
        "extract_caption": "이미지 또는 PDF를 업로드하세요. 추출하고 번역해 드립니다.",
        "extract_button": "텍스트 추출 및 번역",
        "target_language_label": "번역할 언어",
        "file_uploader_label": "클릭해 업로드하거나 파일을 드래그 앤 드롭",
        "file_uploader_hint": "제한 200MB / JPG, JPEG, PNG, PDF",
        "learn_inquire_header": "학습 & 문의",
        "learn_inquire_text": "번역 후, 특정 문장을 선택해 질문할 수 있어요.",
        "history_header": "히스토리",
        "history_text": "최근 번역 기록이 여기에 표시됩니다.",
        "history_item_prefix": "기록",
        "context_label": "문맥",
        "chat_input_label": "질문할 문장/문단 (선택 또는 붙여넣기):",
        "ask_ai_button": "AI에게 질문",
        "error_api_key": "API 키가 설정되지 않아 기능을 실행할 수 없습니다.",
        "error_ocr_fail": "OCR 및 번역에 실패했습니다.",
        "error_api": "Gemini API 오류:",
        "header_korean": "✅ OCR 결과 (원본 한국어)",
        "header_target": "💬 번역 결과 ({target_lang_name})",
        "spinner": "이미지에서 텍스트를 추출하고 {target_lang_name}(으)로 번역 중...",
        "image_header": "🖼️ 업로드된 이미지",
        "error_file_proc": "파일 처리 중 오류가 발생했습니다:",
        "please_upload_first": "먼저 파일을 업로드해 주세요.",
        "pdf_supported": "PDF가 감지되었습니다. 페이지를 선택하세요.",
        "ocr_confidence": "OCR 신뢰도",
        "copy": "복사",
        "save_edits": "수정 내용 저장",
        "export": "내보내기",
        "export_txt": "TXT로 내보내기",
        "export_docx": "DOCX로 내보내기",
        "export_csv": "CSV로 내보내기",
        "original": "원문",
        "translation": "번역문",
        "select_sentences": "문장 선택 (선택된 문장으로 질문)",
        "selected_text_will_be_used": "선택된 문장이 질문에 사용됩니다.",
        "pages": "페이지",
        "saved": "저장되었습니다.",
    },
    "en": {
        "title_main": "ScanTranslate: Korean → Filipino OCR Tool",
        "title_sub": "Extract/translate text from images/PDFs, edit, copy, and export.",
        "app_language": "App Display Language",
        "extract_header": "Extract & Translate Text",
        "extract_caption": "Upload an image or PDF. We’ll extract and translate it.",
        "extract_button": "Extract & Translate Text",
        "target_language_label": "Translate to",
        "file_uploader_label": "Click to upload or drag & drop",
        "file_uploader_hint": "Limit 200MB / JPG, JPEG, PNG, PDF",
        "learn_inquire_header": "Learn & Inquire",
        "learn_inquire_text": "After translating, select specific sentences and ask deeper questions.",
        "history_header": "History",
        "history_text": "Your recent translations will appear here.",
        "history_item_prefix": "History",
        "context_label": "Context",
        "chat_input_label": "Sentence/paragraph to ask about (select or paste):",
        "ask_ai_button": "Ask AI",
        "error_api_key": "API key not set.",
        "error_ocr_fail": "OCR and translation failed.",
        "error_api": "Gemini API Error:",
        "header_korean": "✅ OCR Result (Original Korean)",
        "header_target": "💬 Translation Result ({target_lang_name})",
        "spinner": "Extracting and translating to {target_lang_name}...",
        "image_header": "🖼️ Uploaded Image",
        "error_file_proc": "Error while processing the file:",
        "please_upload_first": "Please upload a file first.",
        "pdf_supported": "PDF detected. Pick a page.",
        "ocr_confidence": "OCR Confidence",
        "copy": "Copy",
        "save_edits": "Save edits",
        "export": "Export",
        "export_txt": "Export TXT",
        "export_docx": "Export DOCX",
        "export_csv": "Export CSV",
        "original": "Original",
        "translation": "Translation",
        "select_sentences": "Select sentences (will be used for the question)",
        "selected_text_will_be_used": "Selected sentences will be used for the question.",
        "pages": "Pages",
        "saved": "Saved.",
    },
    "fil": {
        "title_main": "ScanTranslate: Korean → Filipino OCR Kagamitan",
        "title_sub": "Mag-extract/magsalin mula sa larawan/PDF, mag-edit, kopyahin, at i-export.",
        "app_language": "Wika ng App",
        "extract_header": "I-extract at Isalin ang Teksto",
        "extract_caption": "Mag-upload ng larawan o PDF. I-e-extract at isasalin namin.",
        "extract_button": "I-extract at Isalin",
        "target_language_label": "Isalin sa",
        "file_uploader_label": "I-click para mag-upload o i-drag & drop",
        "file_uploader_hint": "Hangganan 200MB / JPG, JPEG, PNG, PDF",
        "learn_inquire_header": "Matuto at Magtanong",
        "learn_inquire_text": "Pagkatapos magsalin, pumili ng mga pangungusap at magtanong nang mas malalim.",
        "history_header": "Kasaysayan",
        "history_text": "Dito lalabas ang iyong mga huling pagsasalin.",
        "history_item_prefix": "Kasaysayan",
        "context_label": "Konteksto",
        "chat_input_label": "Pangungusap/talata para tanungin (pumili o i-paste):",
        "ask_ai_button": "Itanong sa AI",
        "error_api_key": "Walang API key.",
        "error_ocr_fail": "Bigo ang OCR at pagsasalin.",
        "error_api": "Error sa Gemini API:",
        "header_korean": "✅ Resulta ng OCR (Orihinal na Korean)",
        "header_target": "💬 Resulta ng Pagsasalin ({target_lang_name})",
        "spinner": "Kumukuha at isinasalin sa {target_lang_name}...",
        "image_header": "🖼️ Na-upload na Larawan",
        "error_file_proc": "Error habang pinoproseso ang file:",
        "please_upload_first": "Mag-upload muna ng file.",
        "pdf_supported": "Natukoy ang PDF. Pumili ng pahina.",
        "ocr_confidence": "Kumpiyansa ng OCR",
        "copy": "Kopyahin",
        "save_edits": "I-save ang mga pagbabago",
        "export": "I-export",
        "export_txt": "I-export bilang TXT",
        "export_docx": "I-export bilang DOCX",
        "export_csv": "I-export bilang CSV",
        "original": "Orihinal",
        "translation": "Salin",
        "select_sentences": "Pumili ng mga pangungusap (gagamitin sa tanong)",
        "selected_text_will_be_used": "Gagamitin sa tanong ang napiling pangungusap.",
        "pages": "Mga Pahina",
        "saved": "Nasave.",
    },
}

TARGET_LANGUAGES = {
    "ko": {"code": "Korean", "flag": "🇰🇷", "display_ko": "한국어", "display_en": "Korean", "display_fil": "Koreano"},
    "en": {"code": "English", "flag": "🇺🇸", "display_ko": "영어", "display_en": "English", "display_fil": "Ingles"},
    "fil": {"code": "Filipino (Tagalog)", "flag": "🇵🇭", "display_ko": "필리핀어", "display_en": "Filipino", "display_fil": "Filipino"},
}

# ─────────────────────────────────────────────────────────────
# STATE
# ─────────────────────────────────────────────────────────────
ss = st.session_state
ss.setdefault("app_lang_key", "ko")
ss.setdefault("target_lang_key", "fil")
ss.setdefault("chat_history", [])
ss.setdefault("translation_context", None)
ss.setdefault("history_list", [])
ss.setdefault("edited_korean", "")
ss.setdefault("edited_target", "")
ss.setdefault("ocr_confidence", None)
ss.setdefault("pdf_page_index", 0)

# ─────────────────────────────────────────────────────────────
# STYLE (wide + blue + BIG TITLE)
# ─────────────────────────────────────────────────────────────
st.markdown("""
<style>
.stApp { background-color:#e3f2fd; }
.stApp > header { visibility:hidden; }
div.block-container { max-width:100%; padding: 2rem; }
.center { text-align:center; }
.title-banner {
  background: linear-gradient(90deg,#1976d2,#0d47a1);
  color:#fff;
  padding:14px 30px;
  border-radius:28px;
  font-weight:900;
  display:inline-block;
  font-size: 2.6rem;
  letter-spacing: .3px;
}
.stButton>button { background:#0d47a1 !important; color:#fff !important; border-radius:8px !important; }
.copy-btn { background:#eaf2ff; border:1px solid #98b6ff; padding:6px 10px; border-radius:8px; cursor:pointer; }
.conf-chip{display:inline-block;padding:4px 10px;border-radius:16px;background:#e8f4fd;border:1px solid #9ac3f9;font-weight:600;}
.thumb { border:2px solid #d6e4ff;border-radius:8px; overflow:hidden; }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────
# HELPERS (parsing; caching; copy; export)
# ─────────────────────────────────────────────────────────────
def ui_text(key): return TEXTS[ss["app_lang_key"]][key]

def label_for(code):
    data = TARGET_LANGUAGES[code]
    k = f"display_{ss['app_lang_key']}"
    return f"{data['flag']} {data.get(k, data['display_en'])}"

def _clean_code_fence(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^```[a-zA-Z0-9_-]*\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    return text.strip()

def _extract_json_block(text: str):
    cleaned = _clean_code_fence(text)
    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start != -1 and end != -1 and end > start:
        return cleaned[start:end+1]
    return None

def _heuristic_split(raw: str) -> tuple[str, str]:
    cleaned = _clean_code_fence(raw)
    m_k = re.search(r'"?korean"?\s*:\s*"?(.*?)"?\s*(?:,|\n|$)', cleaned, flags=re.S|re.I)
    m_t = re.search(r'"?(translation|filipino|english)"?\s*:\s*"?(.*?)"?\s*(?:,|\n|$)', cleaned, flags=re.S|re.I)
    if m_k and m_t:
        return (m_k.group(1).strip(), m_t.group(2).strip())
    m1 = re.search(r"원본\(한국어\)\s*:?\s*(.+?)\n+\s*번역\(.+?\)\s*:\s*(.+)$", cleaned, flags=re.S)
    if m1:
        return (m1.group(1).strip(), m1.group(2).strip())
    return (cleaned.strip(), "")

def components_copy_button(uid: str, text: str, label: str):
    import streamlit.components.v1 as components
    html = f"""
    <div>
      <button class="copy-btn" id="btn-{uid}">{label}</button>
      <span id="ok-{uid}" style="margin-left:6px;font-size:.85rem;color:#0d47a1;display:none;">✓</span>
    </div>
    <script>
      const btn = document.getElementById("btn-{uid}");
      const ok = document.getElementById("ok-{uid}");
      btn.onclick = async () => {{
        await navigator.clipboard.writeText({json.dumps(text or "")});
        ok.style.display = "inline";
        setTimeout(()=>{{ ok.style.display="none"; }}, 1200);
      }};
    </script>
    """
    components.html(html, height=36)

def export_docx(korean_text, target_text) -> bytes:
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    doc.add_heading('ScanTranslate Export', level=1)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_heading(ui_text("original"), level=2)
    doc.add_paragraph(korean_text or "")
    doc.add_heading(ui_text("translation"), level=2)
    doc.add_paragraph(target_text or "")
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def export_csv(korean_text, target_text) -> bytes:
    df = pd.DataFrame([{"original": korean_text, "translation": target_text}])
    return df.to_csv(index=False).encode("utf-8-sig")

def sentences_of(text):
    chunks = re.split(r'(?<=[.!?])\s+', (text or "").strip())
    return [s for s in chunks if s]

def _hash_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()

# ─────────────────────────────────────────────────────────────
# NEW: Learn & Inquire helper  ✅
# ─────────────────────────────────────────────────────────────
def generate_inquiry_response(client, question: str, context, focus_text: str = "", model: str = "gemini-2.0-flash") -> str:
    if not question or not question.strip():
        return "질문을 입력해주세요. (예: 이 문장의 의미를 쉽게 설명해 주세요.)"

    # Build a readable context block even if a dict is passed
    if isinstance(context, dict):
        context_block = (
            f"[Korean]\n{context.get('korean','')}\n\n"
            f"[Translation]\n{context.get('target','')}\n\n"
            f"[Target language]\n{context.get('lang','')}"
        )
    else:
        context_block = str(context or "")

    prompt = "\n".join([
        "You are a precise bilingual explainer. Answer briefly but clearly.",
        "",
        "=== Focused Text ===",
        (focus_text or "[None selected]"),
        "",
        "=== Full Context ===",
        context_block,
        "",
        "=== User Question ===",
        question.strip()
    ])

    if client is None:
        return "Gemini 클라이언트가 초기화되지 않았습니다. GEMINI_API_KEY를 확인하세요."

    try:
        resp = client.models.generate_content(
            model=model,
            contents=prompt,
        )
        return (resp.text or "").strip() or "No answer generated."
    except APIError as e:
        return f"AI error: {e}"
    except Exception as e:
        return f"Unexpected error while asking AI: {e}"


# ─────────────────────────────────────────────────────────────
# CACHED: OCR + Translation  (keyed by image bytes hash + target)
# ─────────────────────────────────────────────────────────────
@st.cache_data(show_spinner=False)
def ocr_translate_cached(image_bytes: bytes, mime_type: str, target_lang_name: str, app_lang_key: str):
    """Cache the full OCR+translation result by content + language."""
    if not client:
        return TEXTS[app_lang_key]['error_api_key'], "", None

    image_part = Part.from_bytes(data=image_bytes, mime_type=mime_type)
    prompt = (
        "Perform OCR on the image (Korean expected) and translate to "
        f"{target_lang_name}. Return STRICT JSON ONLY with keys: "
        '{"korean":"...", "translation":"...", "confidence": 0-100}. '
        "Do not add markdown/code fences. Preserve line breaks in 'korean'."
    )
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[prompt, image_part],
        )
        raw = (response.text or "").strip()

        json_block = _extract_json_block(raw)
        korean_result, target_result, conf = "", "", None

        if json_block:
            try:
                data = json.loads(json_block)
                korean_result = (data.get("korean") or "").strip()
                target_result = (data.get("translation") or "").strip()
                conf_val = data.get("confidence", None)
                try:
                    conf = int(round(float(conf_val))) if conf_val is not None else None
                except Exception:
                    conf = None
            except Exception:
                korean_result, target_result = _heuristic_split(raw)
                conf = None
        else:
            korean_result, target_result = _heuristic_split(raw)
            conf = None

        # Keep boxes separate
        if not target_result:
            target_result = ""
        if not korean_result:
            korean_result = ""

        return korean_result, target_result, conf

    except APIError as e:
        return f"{TEXTS[app_lang_key]['error_api']} {e}", "", None
    except Exception as e:
        return f"{TEXTS[app_lang_key]['error_ocr_fail']} 오류: {e}", "", None

# ─────────────────────────────────────────────────────────────
# CACHED: Render a single PDF page thumbnail (lazy)
# ─────────────────────────────────────────────────────────────
@st.cache_data(show_spinner=False)
def render_pdf_page_thumb(pdf_bytes: bytes, page_index: int, scale: float = 1.2) -> bytes:
    """
    Render one page as PNG bytes (lighter scale keeps it fast).
    Cached by (pdf hash + page index + scale).
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    p = doc.load_page(page_index)
    pix = p.get_pixmap(matrix=fitz.Matrix(scale, scale))
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    bio = io.BytesIO()
    img.save(bio, format="PNG")
    return bio.getvalue()

# ─────────────────────────────────────────────────────────────
# HEADER
# ─────────────────────────────────────────────────────────────
st.markdown('<div class="center"><span class="title-banner">'
            f'{TEXTS[ss["app_lang_key"]]["title_main"]}'
            '</span></div>', unsafe_allow_html=True)
st.markdown(f'<p class="center" style="color:#344054">{TEXTS[ss["app_lang_key"]]["title_sub"]}</p>', unsafe_allow_html=True)

# App language selector
app_lang_options = {"ko": "🇰🇷", "en": "🇺🇸", "fil": "🇵🇭"}
col_lang, _ = st.columns([1,8])
with col_lang:
    st.markdown(f"**🌐 {TEXTS[ss['app_lang_key']]['app_language']}**")
    chosen = st.selectbox(
        "app_lang",
        list(app_lang_options.keys()),
        format_func=lambda x: app_lang_options[x],
        label_visibility="collapsed",
        index=list(app_lang_options.keys()).index(ss["app_lang_key"])
    )
    if chosen != ss["app_lang_key"]:
        ss["app_lang_key"] = chosen
        st.rerun()

# ─────────────────────────────────────────────────────────────
# LAYOUT COLUMNS
# ─────────────────────────────────────────────────────────────
left, right = st.columns([2,1], gap="large")

with left:
    # Upload + controls (wrapped in a form to avoid extra reruns)
    with st.container(border=True):
        st.subheader(ui_text("extract_header"))
        st.caption(ui_text("extract_caption"))

        with st.form("extract_form", clear_on_submit=False):
            st.markdown(
                f"<div class='center' style='color:#1976d2'><b>{ui_text('file_uploader_label')}</b>"
                f"<br><small>{ui_text('file_uploader_hint')}</small></div>",
                unsafe_allow_html=True
            )
            uploaded = st.file_uploader("upload", type=["jpg","jpeg","png","pdf"], label_visibility="collapsed")

            lc, bc = st.columns([1.5,2])
            with lc:
                st.markdown(f"**💬 {ui_text('target_language_label')}**")
                chosen_tgt = st.selectbox(
                    "tgt",
                    list(TARGET_LANGUAGES.keys()),
                    format_func=label_for,
                    label_visibility="collapsed",
                    index=list(TARGET_LANGUAGES.keys()).index(ss["target_lang_key"])
                )
            with bc:
                submitted = st.form_submit_button(ui_text("extract_button"), use_container_width=True)

        ss["target_lang_key"] = chosen_tgt

    # PDF controls
    selected_image_bytes, selected_mime = None, None
    pdf_bytes = None
    if uploaded is not None and uploaded.type == "application/pdf":
        st.info(ui_text("pdf_supported"))
        pdf_bytes = uploaded.read()
        # quick page count (open once)
        _doc_tmp = fitz.open(stream=pdf_bytes, filetype="pdf")
        page_count = _doc_tmp.page_count

        # Thumbnails (lazy render): show up to 6 thumbs
        show_n = min(6, page_count)
        cols = st.columns(show_n) if show_n else []
        for i in range(show_n):
            thumb_png = render_pdf_page_thumb(pdf_bytes, i, 1.0)  # lighter & cached
            with cols[i]:
<<<<<<< HEAD
                st.image(thumb_png, caption=f"{ui_text('pages')} {i+1}", use_container_width=True)
=======
                st.image(thumb_png, caption=f"{ui_text('pages')} {i+1}", use_column_width=True)
>>>>>>> ea88045 (Initial commit: add ScanTranslate project)

        # Picker
        ss["pdf_page_index"] = st.slider(f"{ui_text('pages')}", 1, page_count, ss.get("pdf_page_index", 0) + 1) - 1
        # Render the selected page to feed OCR (cached)
        selected_image_bytes = render_pdf_page_thumb(pdf_bytes, ss["pdf_page_index"], 1.4)
        selected_mime = "image/png"

    # PROCESS (only when form submitted)
    if uploaded is not None and 'submitted' in locals() and submitted:
        with st.container(border=True):
            try:
                if uploaded.type in ["image/jpeg","image/png","image/jpg"]:
                    image = Image.open(uploaded)
                    img_bio = io.BytesIO()
                    fmt = uploaded.type.split('/')[-1]
                    image.save(img_bio, format=fmt)
                    image_bytes = img_bio.getvalue()
                    mime = uploaded.type
                elif uploaded.type == "application/pdf":
                    image_bytes = selected_image_bytes
                    mime = selected_mime
                else:
                    st.error("Unsupported file.")
                    image_bytes, mime = None, None

                if image_bytes:
                    spinner_text = ui_text("spinner").format(
                        target_lang_name=TARGET_LANGUAGES[ss['target_lang_key']]['code']
                    )
                    with st.spinner(spinner_text):
                        # cache by content hash + target
                        key_hash = _hash_bytes(image_bytes)
                        korean_result, target_result, conf = ocr_translate_cached(
                            image_bytes=image_bytes,
                            mime_type=mime,
                            target_lang_name=TARGET_LANGUAGES[ss['target_lang_key']]['code'],
                            app_lang_key=ss["app_lang_key"],
                        )

                    # Save strictly separated content
                    ss["edited_korean"] = korean_result or ""
                    ss["edited_target"] = target_result or ""
                    ss["ocr_confidence"] = conf
                    ss["translation_context"] = {
                        "korean": ss["edited_korean"],
                        "target": ss["edited_target"],
                        "lang": TARGET_LANGUAGES[ss['target_lang_key']]["code"]
                    }
                    new_hist = {
                        "korean": ss["edited_korean"],
                        "target": ss["edited_target"],
                        "lang_name": TARGET_LANGUAGES[ss['target_lang_key']]["code"],
                        "lang_flag": TARGET_LANGUAGES[ss['target_lang_key']]["flag"],
                        "confidence": conf
                    }
                    ss["history_list"].insert(0, new_hist)
                    ss["history_list"] = ss["history_list"][:5]

            except Exception as e:
                st.error(f"{ui_text('error_file_proc')} {e}")

    # Side-by-side editor + copy + export
    if ss.get("edited_korean") or ss.get("edited_target"):
        st.markdown("### ✍️ Side-by-Side Editor")
        c1, c2 = st.columns(2)
        with c1:
            conf_html = (f"<span class='conf-chip'>{ui_text('ocr_confidence')}: {ss['ocr_confidence']}%</span>"
                         if ss['ocr_confidence'] is not None else "")
            st.markdown(f"**{ui_text('original')}**  {conf_html}", unsafe_allow_html=True)
            ss["edited_korean"] = st.text_area("kor", ss["edited_korean"], height=220, label_visibility="collapsed")
            components_copy_button("korean", ss["edited_korean"], ui_text("copy"))
        with c2:
            st.markdown(f"**{ui_text('translation')}**")
            ss["edited_target"] = st.text_area("tgt", ss["edited_target"], height=220, label_visibility="collapsed")
            components_copy_button("target", ss["edited_target"], ui_text("copy"))

        sc, ec1, ec2, ec3 = st.columns([1,1,1,1])
        with sc:
            if st.button(ui_text("save_edits"), use_container_width=True):
                if ss["history_list"]:
                    ss["history_list"][0]["korean"] = ss["edited_korean"]
                    ss["history_list"][0]["target"] = ss["edited_target"]
                ss["translation_context"] = {
                    "korean": ss["edited_korean"],
                    "target": ss["edited_target"],
                    "lang": TARGET_LANGUAGES[ss['target_lang_key']]["code"]
                }
                st.success(ui_text("saved"))
        with ec1:
            st.download_button(
                ui_text("export_txt"),
                data=(ss["edited_korean"]+"\n\n---\n\n"+ss["edited_target"]).encode("utf-8"),
                file_name=f"scantranslate_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                mime="text/plain",
                use_container_width=True
            )
        with ec2:
            st.download_button(
                ui_text("export_docx"),
                data=export_docx(ss["edited_korean"], ss["edited_target"]),
                file_name=f"scantranslate_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        with ec3:
            st.download_button(
                ui_text("export_csv"),
                data=export_csv(ss["edited_korean"], ss["edited_target"]),
                file_name=f"scantranslate_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                mime="text/csv",
                use_container_width=True
            )

with right:
    # Learn & Inquire (form to avoid reruns while typing)
    with st.container(border=True):
        st.subheader(ui_text("learn_inquire_header"))
        if ss.get("translation_context"):
            ctx = ss["translation_context"]
            st.caption(ui_text("learn_inquire_text"))
            st.markdown("---")

            sents = sentences_of(ss["edited_target"])
            with st.form("ask_ai_form", clear_on_submit=False):
                selected = st.multiselect(
                    ui_text("select_sentences"),
                    options=[f"{i+1}. {s}" for i, s in enumerate(sents)],
                    help=ui_text("selected_text_will_be_used")
                )
                # This textarea is the actual QUESTION
                pasted_question = st.text_area(ui_text("chat_input_label"), height=100)
                ask = st.form_submit_button(ui_text("ask_ai_button"), use_container_width=True)

            if ask:
                # Build focus text from selections (strip the "1. " prefixes)
                focus_text = " ".join([opt.split(". ", 1)[-1] for opt in selected]) if selected else ""
                question = pasted_question.strip()

                if not question:
                    st.warning("질문을 입력해주세요. 예: '이 문장의 의미를 쉽게 설명해 줘.'")
                else:
                    with st.spinner("..."):
                        # ✅ Correct signature (client first) and ctx can be dict
                        answer = generate_inquiry_response(client, question, ctx, focus_text=focus_text)
                    ss["chat_history"].append(("user", question))
                    ss["chat_history"].append(("model", answer))
                    st.markdown(f"**👤 User:** *{question}*")
                    st.markdown(f"**🤖 AI Tutor:** {answer}")

            if ss["chat_history"]:
                st.markdown("---")
                for role, text in ss["chat_history"]:
                    if role == "user":
                        st.markdown(f"**👤 User:** *{text}*")
                    else:
                        st.markdown(f"**🤖 AI Tutor:** {text}")
        else:
            st.write(ui_text("history_text"))

    st.markdown("<br>", unsafe_allow_html=True)

    # History
    with st.container(border=True):
        st.subheader(f"◷ {ui_text('history_header')}")
        if ss["history_list"]:
            for i, entry in enumerate(ss["history_list"]):
                label = f"{ui_text('history_item_prefix')} #{i+1}: Korean → {entry['lang_flag']} {entry['lang_name']}"
                with st.expander(label):
                    if entry.get("confidence") is not None:
                        st.caption(f"{ui_text('ocr_confidence')}: {entry['confidence']}%")
                    st.caption(ui_text("original"))
                    st.code(entry["korean"])
                    st.caption(ui_text("translation"))
                    st.code(entry["target"])
        else:
            st.write(ui_text("history_text"))
